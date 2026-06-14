"""把论文markdown转为Word (.docx)"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font; font.name = '宋体'; font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_title(text, level=0):
    if level == 0:  # 主标题
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text); run.bold = True; run.font.size = Pt(16); run.font.name = '黑体'
        p.space_after = Pt(6)
    elif level == 1:  # 一级标题
        p = doc.add_paragraph(); run = p.add_run(text); run.bold = True; run.font.size = Pt(14)
        p.space_before = Pt(12); p.space_after = Pt(6)
    elif level == 2:  # 二级标题
        p = doc.add_paragraph(); run = p.add_run(text); run.bold = True; run.font.size = Pt(12)
        p.space_before = Pt(6); p.space_after = Pt(3)

def add_body(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold→plain
    text = re.sub(r'`(.+?)`', r'\1', text)         # code→plain
    text = re.sub(r'\$([^$]+)\$', r'\1', text)     # inline math→plain
    text = re.sub(r'\$\$.*?\$\$', '', text, flags=re.DOTALL)  # block math→remove
    text = re.sub(r'\[(\d+)\]', r'[\1]', text)      # citation→plain
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)     # image→remove
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # link→text
    text = re.sub(r'\|.*?\|', '', text)             # table→remove
    text = re.sub(r'[-]{3,}', '', text)             # hr→remove
    if text.strip():
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
        p.paragraph_format.line_spacing = 1.25

def add_table_row(cells):
    row = doc.add_table(rows=1, cols=len(cells)).rows[0]
    for i, c in enumerate(cells):
        row.cells[i].text = c.strip()
        for p in row.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.font.size = Pt(10)

# 读取markdown
with open('paper_tangdun.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_code = in_table = in_math = False
for line in lines:
    stripped = line.strip()

    if stripped.startswith('$$'): in_math = not in_math; continue
    if in_math: continue

    if stripped.startswith('# ') or stripped.startswith('# 基于'):
        add_title(stripped.lstrip('# '), 0)
    elif stripped.startswith('## '):
        add_title(stripped.lstrip('## '), 1)
    elif stripped.startswith('### '):
        add_title(stripped.lstrip('### '), 2)
    elif '|' in stripped and stripped.count('|') >= 2:
        continue  # skip tables
    elif stripped == '---': continue
    elif stripped.startswith('> '): continue
    elif stripped.startswith('**关键词'):
        add_body(stripped)
    else:
        add_body(stripped)

output = '糖盾系统学术论文.docx'
doc.save(output)
print(f'已生成: {output}')
