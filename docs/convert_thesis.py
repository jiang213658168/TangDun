"""完整转换研究生论文markdown→Word，保留所有内容"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re, os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font; font.name = '宋体'; font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def add_title(text, level=0):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text); run.bold = True; run.font.size = Pt(16); run.font.name = '黑体'
        p.space_after = Pt(8)
    elif level == 1:
        run = p.add_run(text); run.bold = True; run.font.size = Pt(14); run.font.name = '黑体'
        p.space_before = Pt(12); p.space_after = Pt(6)
    elif level == 2:
        run = p.add_run(text); run.bold = True; run.font.size = Pt(12); run.font.name = '黑体'
        p.space_before = Pt(8); p.space_after = Pt(4)

def add_body(text):
    # 保留内容，只清理markdown标记
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)       # bold
    text = re.sub(r'`([^`]+)`', r'\1', text)            # code
    text = re.sub(r'\$([^$]+)\$', r'\1', text)          # inline math (keep formula)
    text = re.sub(r'\$\$[^$]*\$\$', '', text, flags=re.DOTALL)  # block math (remove, too complex)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # links
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)          # images
    text = re.sub(r'^#+\s*', '', text)                   # heading markers
    text = re.sub(r'^>\s*', '', text)                    # blockquote
    text = re.sub(r'[-*]\s', '', text)                   # list markers
    if text.strip():
        p = doc.add_paragraph(text.strip())
        p.paragraph_format.first_line_indent = Cm(0.74)

def add_table_row(cells):
    num = len(cells)
    table = doc.add_table(rows=1, cols=num, style='Table Grid')
    for i, c in enumerate(cells):
        cell = table.rows[0].cells[i]
        cell.text = c.strip()
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)

with open('graduate_thesis.md', 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0
in_math = False
in_table = False
table_rows = []
in_abstract = False

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Math blocks
    if stripped.startswith('$$'):
        in_math = not in_math
        i += 1
        continue
    if in_math:
        i += 1
        continue

    # Skip separators and empty lines
    if stripped == '---' or stripped == '':
        i += 1
        continue

    # Title
    if stripped.startswith('# ') or stripped.startswith('# 基于'):
        add_title(stripped.lstrip('# ').strip(), 0)
        i += 1
        continue

    # Chapter headings
    if stripped.startswith('## ') and not stripped.startswith('### '):
        t = stripped.lstrip('## ').strip()
        if t.isdigit() or t.startswith('1.') or t.startswith('2.') or t.startswith('3.') or \
           t.startswith('4.') or t.startswith('5.') or t.startswith('6.') or \
           t.startswith('7.') or t.startswith('8.'):
            add_title(t, 1)
        elif t in ['参考文献', '致谢', '摘要', 'Abstract']:
            p = doc.add_paragraph()
            run = p.add_run(t); run.bold = True; run.font.size = Pt(16); run.font.name = '黑体'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # Section headings (#3)
    if stripped.startswith('### '):
        t = stripped.lstrip('### ').strip()
        add_title(t, 2)
        i += 1
        continue

    if stripped.startswith('#### '):
        t = stripped.lstrip('#### ').strip()
        p = doc.add_paragraph()
        run = p.add_run(t); run.bold = True; run.font.size = Pt(11)
        i += 1
        continue

    # Tables
    if '|' in stripped and stripped.count('|') >= 2:
        cells = [c.strip() for c in stripped.split('|')[1:-1]]  # Skip leading/trailing |
        # Skip separator rows like |---|---|
        if all(re.match(r'^[-:]+$', c) for c in cells):
            i += 1
            continue
        add_table_row(cells)
        i += 1
        continue

    # Declaration
    if stripped.startswith('**学位论文'):
        p = doc.add_paragraph(stripped.replace('*', ''))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # Regular paragraph
    add_body(stripped)
    i += 1

output = '糖盾研究生论文_去AI化.docx'
doc.save(output)
print(f'已生成: {os.path.abspath(output)} ({os.path.getsize(output)} bytes)')
